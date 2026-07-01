"""
LegalForge Backend - Full Verification Script
Tests all API endpoints in sequence.
"""
import requests
import json
import os
import sys
import io

# Force UTF-8 output
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')

BASE = "http://localhost:8000"
PASS = "[PASS]"
FAIL = "[FAIL]"
results = []
log_lines = []

def log(msg):
    print(msg)
    log_lines.append(msg)

def test(name, fn):
    try:
        ok, detail = fn()
        status = PASS if ok else FAIL
        results.append((name, ok, detail))
        log(f"  {status} {name}: {detail}")
    except Exception as e:
        results.append((name, False, str(e)))
        log(f"  {FAIL} {name}: EXCEPTION - {e}")


def separator(title):
    log(f"\n{'='*60}")
    log(f"  {title}")
    log(f"{'='*60}")


# ============================================================
# STEP 1: Health / Root
# ============================================================
separator("STEP 1: Health & Root Endpoint")

def test_root():
    r = requests.get(f"{BASE}/")
    data = r.json()
    return r.status_code == 200 and data.get("status") == "online", f"[{r.status_code}] {json.dumps(data)}"

def test_health():
    r = requests.get(f"{BASE}/health")
    data = r.json()
    return r.status_code == 200 and data.get("api") == True, f"[{r.status_code}] {json.dumps(data)}"

test("GET /", test_root)
test("GET /health", test_health)

# ============================================================
# STEP 2: Auth Endpoints
# ============================================================
separator("STEP 2: Auth Endpoints (Register + Login)")

auth_headers = {}


def test_register():
    r = requests.post(f"{BASE}/api/register", json={
        "name": "Test User",
        "email": "test@legalforge.com",
        "password": "TestPass123"
    })
    data = r.json()
    if r.status_code == 200:
        return data.get("success") == True, f"[{r.status_code}] got success={bool(data.get('success'))}"
    elif r.status_code == 400 and "already" in str(data.get("detail", "")).lower():
        return True, f"[{r.status_code}] User already exists (OK for re-run)"
    return False, f"[{r.status_code}] {json.dumps(data)}"

def test_login():
    global auth_headers
    r = requests.post(f"{BASE}/api/login", json={
        "email": "test@legalforge.com",
        "password": "TestPass123"
    })
    data = r.json()
    if r.status_code == 200 and "access_token" in data:
        auth_headers = {"Authorization": f"Bearer {data['access_token']}"}
        return True, f"[{r.status_code}] got token=True"
    return False, f"[{r.status_code}] {json.dumps(data)}"

def test_login_bad():
    r = requests.post(f"{BASE}/api/login", json={
        "email": "wrong@email.com",
        "password": "wrongpass"
    })
    return r.status_code == 401, f"[{r.status_code}] correctly rejected bad creds"

test("POST /api/register", test_register)
test("POST /api/login (valid)", test_login)
test("POST /api/login (invalid)", test_login_bad)

# ============================================================
# STEP 3: Legal Search
# ============================================================
separator("STEP 3: Legal Search (RAG / Sentence Transformers)")

def test_legal_search():
    r = requests.post(f"{BASE}/api/legal-search", json={
        "query": "breach of contract remedies in commercial agreements"
    })
    data = r.json()
    return r.status_code == 200, f"[{r.status_code}] keys={list(data.keys())}"

def test_legal_search_short():
    r = requests.post(f"{BASE}/api/legal-search", json={
        "query": "ab"
    })
    return r.status_code == 400, f"[{r.status_code}] correctly rejected short query"

test("POST /api/legal-search (valid)", test_legal_search)
test("POST /api/legal-search (too short)", test_legal_search_short)

# ============================================================
# STEP 4: Draft Generation
# ============================================================
separator("STEP 4: Draft Generation + PDF Export")

draft_id = None

def test_generate_draft():
    global draft_id
    r = requests.post(f"{BASE}/api/generate-draft", headers=auth_headers, json={
        "doc_type": "Legal Notice",
        "client_name": "Rahul Sharma",
        "opposing_party": "ABC Corp Pvt Ltd",
        "case_description": "The opposing party has failed to pay outstanding dues of Rs. 5,00,000 for software development services rendered between January 2025 and June 2025."
    })
    data = r.json()
    if r.status_code == 200 and "draft_id" in data:
        draft_id = data["draft_id"]
        return True, f"[{r.status_code}] draft_id={draft_id}, has preview={bool(data.get('preview_text'))}"
    return False, f"[{r.status_code}] {json.dumps(data)}"

def test_generate_consumer_complaint():
    r = requests.post(f"{BASE}/api/generate-draft", headers=auth_headers, json={
        "doc_type": "Consumer Complaint",
        "client_name": "Priya Mehta",
        "opposing_party": "XYZ Electronics",
        "case_description": "Purchased a laptop that stopped working within 2 weeks. The company refused to honor the warranty."
    })
    data = r.json()
    return r.status_code == 200 and "draft_id" in data, f"[{r.status_code}] type=Consumer Complaint"

def test_generate_rental():
    r = requests.post(f"{BASE}/api/generate-draft", headers=auth_headers, json={
        "doc_type": "Rental Agreement",
        "client_name": "Amit Verma",
        "opposing_party": "Sneha Gupta",
        "case_description": "Residential property lease for a 2BHK flat located at MG Road, Bangalore."
    })
    data = r.json()
    return r.status_code == 200 and "draft_id" in data, f"[{r.status_code}] type=Rental Agreement"

def test_generate_invalid_type():
    r = requests.post(f"{BASE}/api/generate-draft", headers=auth_headers, json={
        "doc_type": "Invalid Type",
        "client_name": "Test",
        "opposing_party": "Test",
        "case_description": "Test"
    })
    return r.status_code == 400, f"[{r.status_code}] correctly rejected invalid doc_type"

test("POST /api/generate-draft (Legal Notice)", test_generate_draft)
test("POST /api/generate-draft (Consumer Complaint)", test_generate_consumer_complaint)
test("POST /api/generate-draft (Rental Agreement)", test_generate_rental)
test("POST /api/generate-draft (Invalid type)", test_generate_invalid_type)

# Download draft
def test_download_draft():
    if not draft_id:
        return False, "No draft_id from previous test"
    r = requests.get(f"{BASE}/api/download-draft/{draft_id}")
    return r.status_code == 200 and "application/pdf" in r.headers.get("content-type", ""), \
           f"[{r.status_code}] content-type={r.headers.get('content-type')}, size={len(r.content)} bytes"

def test_download_draft_invalid():
    r = requests.get(f"{BASE}/api/download-draft/nonexistent123")
    return r.status_code == 404, f"[{r.status_code}] correctly returned 404 for missing draft"

test("GET /api/download-draft/{id} (valid)", test_download_draft)
test("GET /api/download-draft/{id} (invalid)", test_download_draft_invalid)

# ============================================================
# STEP 5: Generate fake DOCX & test /api/analyze-contract
# ============================================================
separator("STEP 5: Contract Analysis (DOCX Upload)")

# Create a fake contract DOCX
from docx import Document as DocxDocument

def create_test_contract():
    doc = DocxDocument()
    doc.add_heading("SERVICE AGREEMENT", level=1)
    doc.add_paragraph(f"This Service Agreement is entered into on March 14, 2026.")

    doc.add_heading("1. Parties", level=2)
    doc.add_paragraph("This agreement is between Alpha Technologies Pvt Ltd (\"Service Provider\") and Beta Corp Ltd (\"Client\").")

    doc.add_heading("2. Scope of Services", level=2)
    doc.add_paragraph("The Service Provider shall deliver software development services including design, development, testing, and deployment of a web application as described in Annexure A.")

    doc.add_heading("3. Payment Terms", level=2)
    doc.add_paragraph("The Client shall pay a total fee of Rs. 15,00,000 in three installments. Payment is due within 30 days of invoice. Late payment shall incur interest at 2% per month.")

    doc.add_heading("4. Liability", level=2)
    doc.add_paragraph("The Service Provider assumes unlimited liability for any damages arising from breach of this agreement. The Client may seek damages without limitation at sole discretion of the Client.")

    doc.add_heading("5. Confidentiality", level=2)
    doc.add_paragraph("Both parties agree to maintain strict confidentiality of all proprietary information shared during the course of this engagement. This non-disclosure obligation survives termination.")

    doc.add_heading("6. Termination", level=2)
    doc.add_paragraph("Either party may unilateral terminate this agreement with 15 days written notice. Upon termination, the Client shall pay for all work completed. The contract includes auto-renewal for subsequent 12-month periods unless either party provides 90 day notice.")

    doc.add_heading("7. Intellectual Property", level=2)
    doc.add_paragraph("The Client assigns all rights, title, and interest in the deliverables upon full payment. The Service Provider retains a perpetual license to use the underlying frameworks.")

    doc.add_heading("8. Non-Compete", level=2)
    doc.add_paragraph("The Service Provider agrees to a non-compete clause restricting similar services to competitors for 24 months after termination.")

    doc.add_heading("9. Indemnification", level=2)
    doc.add_paragraph("The Service Provider shall indemnify and hold harmless the Client from any claims, damages, or losses arising from the services. The Client waives all rights to challenge the indemnification terms.")

    doc.add_heading("10. Force Majeure", level=2)
    doc.add_paragraph("Neither party shall be liable for delays caused by force majeure events including but not limited to natural disasters, pandemics, or government actions.")

    doc.add_heading("11. Dispute Resolution", level=2)
    doc.add_paragraph("Any disputes arising under this agreement shall be resolved by mutual agreement. If mutual agreement cannot be reached within 30 days, the matter shall be referred to arbitration under the laws of India.")

    doc.add_heading("12. Governing Law", level=2)  
    doc.add_paragraph("This agreement shall be governed by the applicable law of India, with jurisdiction in the courts of Bangalore.")

    doc.add_heading("13. No Refund Policy", level=2)
    doc.add_paragraph("All payments made under this agreement are non-refundable. No refund shall be issued under any circumstances including early termination or dissatisfaction with deliverables.")

    test_path = os.path.join(os.path.dirname(__file__), "test_contract_verify.docx")
    doc.save(test_path)
    return test_path

docx_path = create_test_contract()
log(f"  Created test contract: {docx_path}")

def test_analyze_contract():
    with open(docx_path, "rb") as f:
        r = requests.post(f"{BASE}/api/analyze-contract", files={"file": ("test_contract.docx", f, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")})
    data = r.json()
    if r.status_code == 200:
        return True, f"[{r.status_code}] risk_score={data.get('risk_score')}, total_clauses={data.get('total_clauses')}, warnings={len(data.get('clause_warnings', []))}"
    return False, f"[{r.status_code}] {json.dumps(data)}"

def test_analyze_contract_bad_ext():
    import tempfile
    tmp = os.path.join(tempfile.gettempdir(), "test.zip")
    with open(tmp, "w") as f:
        f.write("This is a fake zip file.")
    with open(tmp, "rb") as f:
        r = requests.post(f"{BASE}/api/analyze-contract", files={"file": ("test.zip", f, "application/zip")})
    return r.status_code == 400, f"[{r.status_code}] correctly rejected .zip file"

test("POST /api/analyze-contract (DOCX)", test_analyze_contract)
test("POST /api/analyze-contract (bad ext)", test_analyze_contract_bad_ext)

# ============================================================
# STEP 6: Case Insights
# ============================================================
separator("STEP 6: Case Insights Dashboard Data")

def test_case_insights():
    r = requests.get(f"{BASE}/api/case-insights")
    data = r.json()
    return r.status_code == 200 and "total_documents" in data, f"[{r.status_code}] docs={data.get('total_documents')}, searches={data.get('total_searches')}, drafts={data.get('total_drafts')}"

test("GET /api/case-insights", test_case_insights)

# ============================================================
# SUMMARY
# ============================================================
separator("SUMMARY")
passed = sum(1 for _, ok, _ in results if ok)
total = len(results)
log(f"\n  Results: {passed}/{total} tests passed\n")
for name, ok, detail in results:
    status = PASS if ok else FAIL
    log(f"    {status} {name}")
log("")

# Clean up test file
if os.path.exists(docx_path):
    os.remove(docx_path)
    log(f"  Cleaned up: {docx_path}")

# Write results to file
with open(os.path.join(os.path.dirname(__file__), "test_results.txt"), "w", encoding="utf-8") as f:
    f.write("\n".join(log_lines))
